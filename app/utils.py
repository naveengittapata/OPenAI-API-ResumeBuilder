import os
import re
from groq import Groq
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

client = Groq(api_key=os.getenv('GROQ_API_KEY'))

def clean_response(text):
    text = text.replace('**', '')
    text = text.replace('***', '')
    lines = text.strip().split('\n')
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.lower().startswith('here are') or stripped.lower().startswith('here\'s') or stripped.lower().startswith('below are'):
            continue
        cleaned.append(stripped)
    return '\n'.join(cleaned)

def generate_summary(user_input, years_of_experience):
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": f"You are a professional resume writer. Based on the following job description and {years_of_experience} years of experience, generate at least 10 professional summary bullet points that highlight key qualifications, achievements, and expertise. The points should reflect the seniority level matching {years_of_experience} years of experience.\n\nIMPORTANT RULES:\n- Output ONLY the summary points, one per line\n- Start each line with a dash (-) followed by the point\n- Do NOT include any introductory text, headings, or concluding text\n- Do NOT use any markdown formatting like ** or ***\n\nJob Description:\n{user_input}"}],
        max_tokens=800
    )
    return clean_response(response.choices[0].message.content)

def generate_skills_and_tools(user_input):
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": f"List the skills and tools for the following job description grouped by category. Use EXACTLY this format, one category per line:\nCategory Name\tTool1, Tool2, Tool3\n\nUse these categories: Operating Systems, Languages & Frameworks, Web Servers & Tools, Databases, Cloud & DevOps, Testing Tools, Methodologies. Only include categories that are relevant.\n\nIMPORTANT RULES:\n- Use a TAB character between the category name and the tools list\n- Tools within each category should be comma-separated\n- Do NOT include any introductory text, headings, or concluding text\n- Do NOT use any markdown formatting\n- Output ONLY the category lines, nothing else\n\nJob Description:\n{user_input}"}],
        max_tokens=500
    )
    return clean_response(response.choices[0].message.content)

def generate_professional_experience(client_names, start_dates, end_dates, user_input, years_of_experience):
    from dateutil import parser as date_parser

    # Combine client data and sort by end date (most recent first)
    clients = list(zip(client_names, start_dates, end_dates))
    def parse_date(date_str):
        try:
            return date_parser.parse(date_str, fuzzy=True)
        except (ValueError, TypeError):
            return date_parser.parse("Jan 1900")
    clients.sort(key=lambda c: parse_date(c[2]), reverse=True)

    professional_experience = ""
    for name, start, end in clients:
        responsibilities = generate_responsibilities(user_input, years_of_experience)
        professional_experience += f"Client: {name}\nStart Date: {start}\nEnd Date: {end}\nResponsibilities:\n{responsibilities}\n\n"
    return professional_experience

def generate_responsibilities(user_input, years_of_experience):
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": f"You are a professional resume writer. Based on the following job description and {years_of_experience} years of experience, generate at least 10 detailed, real-time professional responsibilities. The responsibilities should reflect the depth and seniority matching {years_of_experience} years of experience — more senior and strategic for higher experience, more hands-on and learning-focused for fewer years. Each point should be specific, action-oriented, and use strong action verbs.\n\nIMPORTANT RULES:\n- Output ONLY the responsibility points, one per line\n- Start each line with a dash (-) followed by the responsibility\n- Do NOT include any introductory text, headings, or concluding text\n- Do NOT use any markdown formatting like ** or ***\n- Do NOT say things like 'Here are the responsibilities'\n\nJob Description:\n{user_input}"}],
        max_tokens=1000
    )
    return clean_response(response.choices[0].message.content)

def generate_resume(user_input, client_names, start_dates, end_dates, education_details, years_of_experience, user_details):
    summary = generate_summary(user_input, years_of_experience)
    skills_tools = generate_skills_and_tools(user_input)
    professional_experience = generate_professional_experience(client_names, start_dates, end_dates, user_input, years_of_experience)
    education = f"Education: {education_details}"

    resume_content = {
        'user_details': user_details,
        'summary': summary,
        'skills_tools': skills_tools,
        'professional_experience': professional_experience,
        'education': education
    }

    return resume_content

def set_heading_black(heading):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

def save_resume_as_docx_in_memory(resume_content):
    doc = Document()

    # User Details at top-left
    user = resume_content['user_details']
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(user['name'])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    contact_para = doc.add_paragraph()
    email_run = contact_para.add_run(f"{user['email']}  |  {user['phone']}")
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(80, 80, 80)

    # Summary as bullet points
    h = doc.add_heading('Summary', level=1)
    set_heading_black(h)
    summary_lines = resume_content['summary'].split('\n')
    for line in summary_lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('-'):
            doc.add_paragraph(line.lstrip('- ').strip(), style='List Bullet')
        else:
            doc.add_paragraph(line, style='List Bullet')

    # Skills & Tools as a categorized table
    h = doc.add_heading('Skills & Tools', level=1)
    set_heading_black(h)
    skills_text = resume_content['skills_tools']
    skill_rows = []
    for line in skills_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if '\t' in line:
            parts = line.split('\t', 1)
            skill_rows.append((parts[0].strip(), parts[1].strip()))
        elif ':' in line:
            parts = line.split(':', 1)
            skill_rows.append((parts[0].strip(), parts[1].strip()))
        else:
            skill_rows.append(('General', line.strip()))

    if skill_rows:
        table = doc.add_table(rows=len(skill_rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for idx, (category, tools) in enumerate(skill_rows):
            cat_cell = table.cell(idx, 0)
            cat_cell.text = ''
            cat_para = cat_cell.paragraphs[0]
            cat_run = cat_para.add_run(category)
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            cat_run.font.color.rgb = RGBColor(0, 0, 0)

            tools_cell = table.cell(idx, 1)
            tools_cell.text = ''
            tools_para = tools_cell.paragraphs[0]
            tools_run = tools_para.add_run(tools)
            tools_run.font.size = Pt(10)
            tools_run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph()

    # Professional Experience
    h = doc.add_heading('Professional Experience', level=1)
    set_heading_black(h)
    exp_text = resume_content['professional_experience']
    sections = exp_text.strip().split('\n\n')
    for section in sections:
        lines = section.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('Client:') or line.startswith('Start Date:') or line.startswith('End Date:') or line.startswith('Responsibilities:'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.startswith('-'):
                doc.add_paragraph(line.lstrip('- ').strip(), style='List Bullet')
            else:
                doc.add_paragraph(line, style='List Bullet')

    # Education
    h = doc.add_heading('Education', level=1)
    set_heading_black(h)
    doc.add_paragraph(resume_content['education'])

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
